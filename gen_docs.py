"""Generate feature + technical .docx documents for the project."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

desktop = os.path.join(os.environ["USERPROFILE"], "Desktop")

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

# ===================================================================
# FEATURE DOCUMENT
# ===================================================================
fdoc = Document()
fdoc.styles['Normal'].font.size = Pt(11)
fdoc.styles['Normal'].font.name = '微软雅黑'

add_heading_styled(fdoc, "wechat-opencode 功能文档", 0)
fdoc.add_paragraph("版本: 0.2.0 | 测试: 141 passed | Python 3.12 | 22 模块 3000+ 行")
fdoc.add_paragraph("")

add_heading_styled(fdoc, "一、项目概述", 1)
fdoc.add_paragraph("wechat-opencode 是一个手机端远程操控电脑的 AI 智能体工具。通过飞书 Bot，用自然语言指挥电脑——写代码、建项目、搜文件、跑命令、生成PPT、操作浏览器，全部从手机搞定。")

add_heading_styled(fdoc, "二、通信方式", 1)
add_heading_styled(fdoc, "2.1 飞书 Bot (默认)", 2)
fdoc.add_paragraph("创建飞书应用 → 获取 App ID / App Secret → 配置到 config.yaml → 手机飞书给 Bot 发消息操控电脑。官方 API，无封号风险。")
add_heading_styled(fdoc, "2.2 Web 桌面端", 2)
fdoc.add_paragraph("启动项目后自动打开 http://127.0.0.1:8080，macOS 风格聊天界面。手机和 Web 端共享同一个 supervisor session，双向同步。")
add_heading_styled(fdoc, "2.3 微信 (备用)", 2)
fdoc.add_paragraph("通过 wcferry 注入 DLL 连接微信，支持'机器人'联系人模式。需要 WeChat 3.9.12.51，有封号风险。")

add_heading_styled(fdoc, "三、核心功能", 1)
features = [
    ("远程操控电脑", "通过 opencode AI 引擎: 文件读写、Shell 命令、代码编写、网络搜索、Git 操作、注册表编辑、进程管理"),
    ("监工 + 执行双引擎", "监工: 和用户聊天、理解需求、分配任务。执行: 默默干活、汇报进度。两个 AI 分工协作"),
    ("Web 桌面端", "macOS 风格聊天界面，双向同步。实时状态、任务列表、快捷键操作"),
    ("PPT 设计师", "/ppt 主题 → 专员确认参数(内容/风格/配色) → 自动生成精美 PPT。5套专业配色"),
    ("浏览器操控", "Playwright MCP + Python Playwright。打开网页、点击、填表、截图"),
    ("Windows 系统操控", "音量/亮度/蓝牙/通知/打印/电源/壁纸/输入法 —— 全部 PowerShell 一键控制"),
    ("截图发送", "/screen → 手机收到桌面截图"),
    ("文件发送/图片接收", "电脑文件传到手机。手机拍照发给 AI 分析"),
    ("会话管理", "/sessions /1 /2 /3 /new"),
    ("任务追踪", "/tasks /task N /status /cancel"),
    ("规划模式", "/plan 目标 → 自动拆解步骤执行"),
    ("成本统计", "/cost 查看 Token 用量和费用"),
    ("操作回滚", "/undo 撤销上一步(git stash)"),
    ("权限审批", "危险命令需 YES 确认"),
    ("代码 Diff", "执行后自动展示文件变更"),
    ("模型切换", "/model flash/pro 一键切换 DeepSeek"),
    ("热重载", "改 .py 代码 3 秒自动生效"),
    ("MCP 扩展", "Playwright + Filesystem MCP"),
]
for title, desc in features:
    p = fdoc.add_paragraph()
    run = p.add_run(f"\u2022 {title}: ")
    run.bold = True
    p.add_run(desc)

add_heading_styled(fdoc, "四、完整指令表", 1)
commands = [
    ("任意文字", "和监工对话"),
    ("/sessions", "查看执行会话列表"),
    ("/1 /2 /3", "切换到第N个执行会话"),
    ("/new", "重置执行层"),
    ("/plan 目标", "规划并执行目标"),
    ("/tasks", "查看任务列表"),
    ("/task N", "任务详情"),
    ("/status", "当前执行状态"),
    ("/cancel", "取消执行中任务"),
    ("/screen", "截图电脑桌面"),
    ("/cost", "费用统计"),
    ("/model", "查看/切换模型"),
    ("/model flash", "切换 DeepSeek V4 Flash"),
    ("/model pro", "切换 DeepSeek V4 Pro"),
    ("/ppt 主题", "生成精美 PPT"),
    ("/undo", "撤销上一步"),
    ("/help", "查看全部指令"),
    ("/stop", "关闭服务"),
]
table = fdoc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "指令"
hdr[1].text = "功能"
for cmd, desc in commands:
    row = table.add_row().cells
    row[0].text = cmd
    row[1].text = desc

add_heading_styled(fdoc, "五、启动方式", 1)
fdoc.add_paragraph("方式1: 双击 start.bat (推荐)")
fdoc.add_paragraph("方式2: PowerShell 运行 .\\scripts\\run_dev.ps1")
fdoc.add_paragraph("方式3: 命令行 .venv\\Scripts\\python -m wechat_opencode")

add_heading_styled(fdoc, "六、配置", 1)
fdoc.add_paragraph("配置文件: config.yaml")
configs = [
    ("bot_type", "\"wechat\" 或 \"feishu\" (当前: feishu)"),
    ("opencode.project_dir", "工作目录 (当前: C:\\Users\\1)"),
    ("opencode.serve_port", "端口 (当前: 4097)"),
    ("feishu.app_id", "飞书应用 ID"),
    ("feishu.app_secret", "飞书应用密钥"),
    ("DEEPSEEK_API_KEY", "环境变量，DeepSeek API 密钥"),
]
for k, v in configs:
    fdoc.add_paragraph(f"\u2022 {k}: {v}")

fdoc.save(os.path.join(desktop, "wechat-opencode-功能文档.docx"))

# ===================================================================
# TECHNICAL DOCUMENT
# ===================================================================
tdoc = Document()
tdoc.styles['Normal'].font.size = Pt(11)
tdoc.styles['Normal'].font.name = '微软雅黑'

add_heading_styled(tdoc, "wechat-opencode 技术交接文档", 0)
tdoc.add_paragraph("版本: 0.2.0 | 测试: 141 passed | Python 3.12 | 22 模块 3000+ 行")
tdoc.add_paragraph("")

add_heading_styled(tdoc, "一、技术栈", 1)
tdoc.add_paragraph("语言: Python 3.10+")
tdoc.add_paragraph("包管理: pip + setuptools")
tdoc.add_paragraph("测试: pytest (141 用例)")
tdoc.add_paragraph("通信: 飞书 lark-oapi SDK (WebSocket + REST) / 微信 wcferry")
tdoc.add_paragraph("AI 引擎: opencode CLI 1.15.10 (HTTP API)")
tdoc.add_paragraph("AI 模型: DeepSeek (deepseek-chat = V4 Flash 非思考)")
tdoc.add_paragraph("Web 框架: Flask 3.1")
tdoc.add_paragraph("MCP: Playwright + Filesystem")
tdoc.add_paragraph("配置: YAML | 存储: JSON")

add_heading_styled(tdoc, "二、项目结构", 1)
modules = [
    ("__main__.py", "CLI 入口"),
    ("types.py", "数据类、WorkerState、协议标签"),
    ("config.py", "YAML 配置加载 (双模式)"),
    ("core.py", "主编排器 ★ 核心 (570+ 行)"),
    ("session.py", "opencode serve + HTTP API (异步+轮询)"),
    ("feishu_bot.py", "飞书 Bot (WS接收+REST发送+图片+文件)"),
    ("bridge.py", "微信桥接 (wcferry)"),
    ("worker.py", "执行层管理器 ★ (异步+轮询+进度)"),
    ("queue.py", "串行执行队列"),
    ("router.py", "消息路由 (/oc 前缀)"),
    ("formatter.py", "输出格式化"),
    ("task_tracker.py", "任务追踪 (JSON 持久化)"),
    ("context.py", "上下文注入"),
    ("permission.py", "权限审批"),
    ("cost_tracker.py", "费用统计 (Token+费用)"),
    ("git_diff.py", "Git 变更展示"),
    ("undo.py", "操作回滚 (git stash)"),
    ("health.py", "健康监控 (微信用)"),
    ("auto_reload.py", "热重载 (文件监控)"),
    ("screenshot.py", "截图工具 (Playwright)"),
    ("ppt_designer.py", "PPT 设计器 (5套主题)"),
    ("web_ui.py", "Web 桌面端 (Flask + HTML)"),
    ("shutdown.py", "信号处理"),
]
table = tdoc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "模块"
hdr[1].text = "功能"
for name, desc in modules:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = desc

add_heading_styled(tdoc, "三、核心架构", 1)

add_heading_styled(tdoc, "3.1 消息流 (飞书模式)", 2)
tdoc.add_paragraph("手机 → FeishuBot(WS) → core._handle_message → supervisor session(HTTP) → opencode → core._handle_result → 协议拦截([TASK:]/[进度:]/[确认:]) → 推回飞书")

add_heading_styled(tdoc, "3.2 监工+执行双引擎", 2)
tdoc.add_paragraph("启动: _init_supervisor() → 创建监工 session + WorkerManager")
tdoc.add_paragraph("分配: 监工回复 [TASK: 目标] → _handle_result 拦截 → worker.start_task() → 异步执行")
tdoc.add_paragraph("进度: Worker 轮询 [进度: ...] → on_inject() → 注入监工 → 监工回复用户")
tdoc.add_paragraph("确认: [确认: 问题 选项] → on_notify() → 手机 → 用户回复 → 自动转发执行层")

add_heading_styled(tdoc, "3.3 HTTP API 调用", 2)
tdoc.add_paragraph("opencode serve 通信全部走 HTTP API:")
add_code(tdoc, "POST /session                创建会话")
add_code(tdoc, "GET  /session                列出会话")
add_code(tdoc, "POST /session/{id}/message   发送消息(阻塞)")
add_code(tdoc, "POST /session/{id}/prompt_async 发送消息(异步)")
add_code(tdoc, "GET  /session/{id}/message   轮询消息")
add_code(tdoc, "POST /session/{id}/abort     中断会话")

add_heading_styled(tdoc, "3.4 协议标签", 2)
tags = [
    ("[TASK: 目标]", "监工→Bridge→执行: 分配任务"),
    ("[进度: 内容]", "执行→Bridge→监工: 汇报进度"),
    ("[确认: 问题 选项]", "执行→Bridge→用户: 需要选择"),
    ("[结果: 成功/失败]", "执行→Bridge→监工: 完成"),
    ("[CANCEL]", "用户/监工→Bridge: 取消"),
]
for tag, desc in tags:
    tdoc.add_paragraph(f"\u2022 {tag}: {desc}")

add_heading_styled(tdoc, "四、Worker 执行流程", 1)
add_code(tdoc, "1. start_task(task) → create_session → execute_async(prompt)")
add_code(tdoc, "2. poll_loop: 每5秒 poll_messages(sid)")
add_code(tdoc, "3. _process_responses:")
add_code(tdoc, "   [进度:] → on_inject() → 注入监工 session")
add_code(tdoc, "   [确认:] → on_notify() → 手机 + on_confirm()")
add_code(tdoc, "   [结果:] → on_inject() → _finish() → 更新历史")

add_heading_styled(tdoc, "五、启动流程", 1)
add_code(tdoc, "1. session.start_serve() → opencode serve (端口4097)")
add_code(tdoc, "2. queue.start() → 消息队列就绪")
add_code(tdoc, "3. feishu_bot.start() → 飞书 WebSocket 连接")
add_code(tdoc, "4. _init_supervisor() → 创建监工+WorkerManager")
add_code(tdoc, "5. web_ui.start_server() → Flask HTTP(8080)")
add_code(tdoc, "6. reloader.start() → 热重载就绪")
add_code(tdoc, "7. _wait_for_shutdown() → 等待信号")

add_heading_styled(tdoc, "六、配置系统", 1)
add_code(tdoc, "config.yaml → Config 数据类")
add_code(tdoc, "  BotType: wechat/feishu")
add_code(tdoc, "  OpenCodeConfig: project_dir, serve_port(4097), command_timeout(300s)")
add_code(tdoc, "  WeChatConfig: prefix(/oc), bot_remark(机器人)")
add_code(tdoc, "  FeishuConfig: app_id, app_secret")
add_code(tdoc, "  ServiceConfig: heartbeat(30s), auto_restart, log_level")
add_code(tdoc, "")
add_code(tdoc, "opencode.json → AI 模型配置")
add_code(tdoc, '  {"model":"deepseek/deepseek-chat","mcp":{...}}')

add_heading_styled(tdoc, "七、依赖", 1)
deps = [
    "wcferry >= 39.5.1.0", "pyyaml >= 6.0", "lark-oapi (飞书 SDK)",
    "requests", "flask", "python-docx", "playwright 1.54.0",
    "pytest >= 8.0", "pytest-mock >= 3.12",
    "opencode CLI 1.15.10", "Git", "DeepSeek API Key",
]
for d in deps:
    tdoc.add_paragraph(f"\u2022 {d}")

add_heading_styled(tdoc, "八、常见问题", 1)
tdoc.add_paragraph("Q: 启动报 HTTP API did not become ready")
tdoc.add_paragraph("A: 端口被占用。等待30秒或 netstat -ano | findstr 4097 查进程")
tdoc.add_paragraph("Q: 飞书收不到回复")
tdoc.add_paragraph("A: 检查 opencode serve (curl http://127.0.0.1:4097/session) 和 DEEPSEEK_API_KEY")
tdoc.add_paragraph("Q: opencode 启动报 5 requests failed")
tdoc.add_paragraph("A: opencode.json 配置有误，只保留 model 字段")
tdoc.add_paragraph("Q: start.bat 双击闪退")
tdoc.add_paragraph("A: PowerShell 运行 .\\scripts\\run_dev.ps1 查看具体错误")

add_heading_styled(tdoc, "九、开发指南", 1)
add_code(tdoc, "运行测试: .venv\\Scripts\\python -m pytest tests/")
add_code(tdoc, "开发启动: .\\scripts\\run_dev.ps1")
add_code(tdoc, "检查配置: .venv\\Scripts\\python -m wechat_opencode --check")
add_code(tdoc, "测试覆盖: 141 用例")

tdoc.save(os.path.join(desktop, "wechat-opencode-技术交接文档.docx"))
print("Done!")
